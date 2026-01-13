from flask import Flask, request, send_file, render_template, session, redirect, url_for, flash, after_this_request
import os
import pandas as pd
import re
from datetime import datetime
import tempfile
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.drawing.image import Image as ExcelImage
from collections import defaultdict
import hashlib
import math
import threading
import uuid
import html
import openai
from openai import RateLimitError, APIError
import time
import random
from dotenv import load_dotenv
from docx import Document
import sys


app = Flask(__name__)

sys.stdout.reconfigure(encoding='utf-8')

# Load environment variables from .env file
load_dotenv()

app.secret_key = os.getenv('FLASK_SECRET_KEY', 'fallback_secret_key')
cache_dir = os.getenv('REPORT_CACHE_DIR', 'report_cache')
os.makedirs(cache_dir, exist_ok=True)
openai_last_request_ts = 0.0
PROMPT_TEMPLATE = (
    "You will receive a WhatsApp chat snippet (part {part}/{total_parts}). "
    "Act as an Industrial/Organizational Psychology expert with a specialization in Human "
    "Resource Management and Management. Analyze the chat data and generate a psychological "
    "analysis report covering emotions, relationships, psychological conditions, and "
    "communication patterns.\n\n"
    "WhatsApp Chat Data:\n"
    "{chat_text}\n\n"
    "Provide a comprehensive analysis of the individuals' psychological states, emotions, "
    "and relationships."
)

# Set the OpenAI API key from the environment variable
openai.api_key = os.getenv("OPENAI_API_KEY")

if not openai.api_key:
    raise ValueError("OpenAI API key not set. Please configure it in the environment variable.")

report_jobs = {}
report_jobs_lock = threading.Lock()


def update_report_job(job_id, **updates):
    with report_jobs_lock:
        report_jobs.setdefault(job_id, {}).update(updates)


def get_report_job(job_id):
    with report_jobs_lock:
        return report_jobs.get(job_id)


def start_report_job(input_path, suffix, use_cache=True, job_id=None):
    if job_id is None:
        job_id = uuid.uuid4().hex
    session['report_job_id'] = job_id
    report_ids = session.get('report_job_ids', [])
    report_ids.append(job_id)
    session['report_job_ids'] = report_ids
    update_report_job(
        job_id,
        status='queued',
        message='Queued for processing.',
        input_file=input_path,
        cached=False
    )

    def process_report():
        update_report_job(job_id, status='running', message='Generating report...')
        def progress_cb(**updates):
            update_report_job(job_id, **updates)
        try:
            if suffix == '.csv':
                df = pd.read_csv(input_path)
            else:
                df = pd.read_excel(input_path)
            if df.empty:
                update_report_job(job_id, status='error', message='Uploaded file is empty.')
                return
            psychological_analysis, from_cache = generate_psychological_report(
                df,
                progress_cb=progress_cb,
                use_cache=use_cache
            )
            if psychological_analysis.startswith("Error:"):
                update_report_job(job_id, status='error', message=psychological_analysis)
                return
            report_path = generate_word_report(psychological_analysis, "psychological_report.docx")
            html_path = generate_html_report(psychological_analysis, "psychological_report.html")
            message = 'Report ready (from cache).' if from_cache else 'Report ready.'
            update_report_job(
                job_id,
                status='done',
                message=message,
                report_file=report_path,
                html_file=html_path,
                cached=from_cache
            )
        except Exception as exc:
            update_report_job(job_id, status='error', message=f"Error: {exc}")

    worker = threading.Thread(target=process_report, daemon=True)
    worker.start()
    return job_id


def parse_whatsapp_chat(file):
    # Accept both desktop and mobile exports with flexible date/time formats.
    desktop_pattern = (
        r'^\[(\d{1,2})[./](\d{1,2})[./](\d{2,4}), '
        r'(\d{1,2}:\d{2})(?::(\d{2}))?\] (.*?)(?::\s)(.*)$'
    )
    mobile_pattern = (
        r'^(\d{1,2})/(\d{1,2})/(\d{2,4}), '
        r'(\d{1,2}:\d{2})(?:\s?(AM|PM))? - (.*)$'
    )
    chat_data = []
    current_message = None

    # Read and decode the file
    for line in file:
        line = line.decode('utf-8').strip()
        if not line:
            continue
        # Strip common directionality marks/BOM that may prefix exported lines.
        line = line.lstrip('\ufeff\u200e\u200f\u202a\u202b\u202c\u202d\u202e\u2066\u2067\u2068\u2069')

        date_time_obj = None
        sender = None
        message = None

        match = re.match(desktop_pattern, line)
        if match:
            day = match.group(1)
            month = match.group(2)
            year = match.group(3)
            time_str = match.group(4)
            seconds = match.group(5) or '00'
            sender = match.group(6)
            message = match.group(7)

            if len(year) == 2:
                year = f"20{year}"

            date_time_str = f'{day}.{month}.{year} {time_str}:{seconds}'
            try:
                date_time_obj = datetime.strptime(date_time_str, '%d.%m.%Y %H:%M:%S')
            except ValueError as ve:
                print(f"Date parsing error: {ve} for line: {line}")
                continue
        else:
            match = re.match(mobile_pattern, line)
            if match:
                month = match.group(1)
                day = match.group(2)
                year = match.group(3)
                time_str = match.group(4)
                am_pm = match.group(5)
                rest = match.group(6)

                if len(year) == 2:
                    year = f"20{year}"

                if am_pm:
                    time_str = f"{time_str} {am_pm}"
                    time_format = '%m/%d/%Y %I:%M %p'
                else:
                    time_format = '%m/%d/%Y %H:%M'

                date_time_str = f'{month}/{day}/{year} {time_str}'
                try:
                    date_time_obj = datetime.strptime(date_time_str, time_format)
                except ValueError as ve:
                    print(f"Date parsing error: {ve} for line: {line}")
                    continue

                sender, sep, message = rest.partition(': ')
                if not sep:
                    sender = 'System'
                    message = rest

        if date_time_obj is not None:
            chat_data.append({
                'Date': date_time_obj.date(),
                'Time': date_time_obj.time(),
                'Datetime': date_time_obj,
                'Sender': sender,
                'Message': message
            })
            current_message = len(chat_data) - 1
        elif current_message is not None:
            # Handle multi-line messages
            chat_data[current_message]['Message'] += '\n' + line

    return pd.DataFrame(chat_data)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('fileup')
        if file and file.filename.endswith('.txt'):
            # Parse the uploaded file
            parsed_chat = parse_whatsapp_chat(file)

            # Check if parsed_chat is empty
            if parsed_chat.empty:
                flash("No messages were parsed. Check the chat export format and try again.", 'error')
                return redirect(url_for('index'))

            # Calculate average reply times for each sender
            avg_reply_times = defaultdict(list)
            previous_sender = None
            previous_time = None

            for i, row in parsed_chat.iterrows():
                current_sender = row['Sender']
                current_time = row['Datetime']

                if previous_sender and previous_sender != current_sender:
                    time_difference = (current_time - previous_time).total_seconds() / 60  # Convert to minutes
                    avg_reply_times[previous_sender].append(time_difference)

                previous_sender = current_sender
                previous_time = current_time

            avg_reply_times = {sender: sum(times) / len(times) for sender, times in avg_reply_times.items() if times}

            # Create a temporary Excel file
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx').name
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                parsed_chat.to_excel(writer, sheet_name='Parsed Chat', index=False)
                message_summary = parsed_chat['Sender'].value_counts().reset_index()
                message_summary.columns = ['Sender', 'Total Messages']
                message_summary.to_excel(writer, sheet_name='Message Summary', index=False)
                avg_reply_df = pd.DataFrame(list(avg_reply_times.items()), columns=['Sender', 'Average Reply Time (mins)'])
                avg_reply_df.to_excel(writer, sheet_name='Avg Reply Time', index=False)

            # Store the file path for download
            session['output_file'] = output_path
            session['file_downloaded'] = False  # Track download status

            flash("File processed successfully.", "success")
            return redirect(url_for('index'))

    prompt_preview = PROMPT_TEMPLATE.format(
        part=1,
        total_parts=1,
        chat_text="{chat_text}"
    )
    return render_template('index.html', prompt_preview=prompt_preview)

@app.route('/download_report', methods=['GET'])
def download_report():
    report_file = session.get('report_file')
    if not report_file or not os.path.exists(report_file):
        flash("No report available to download.", "error")
        return redirect(url_for('index'))

    # Delete the report file after sending it to the user
    @after_this_request
    def remove_report_file(response):
        try:
            os.remove(report_file)
            session.pop('report_file', None)  # Clear the file from the session
        except Exception as e:
            print(f"Error deleting file: {str(e).encode('utf-8', errors='ignore')}")
        return response

    return send_file(report_file, as_attachment=True, download_name="psychological_report.txt")

@app.route('/download_psychological_report', methods=['GET'])
def download_psychological_report():
    report_file = session.get('report_file')
    if not report_file or not os.path.exists(report_file):
        flash("No psychological report available to download.", "error")
        return redirect(url_for('index'))

    # Delete the report file after sending it to the user
    @after_this_request
    def remove_report_file(response):
        try:
            os.remove(report_file)
            session.pop('report_file', None)  # Clear the file from the session
        except Exception as e:
            print(f"Error deleting file: {str(e)}")
        return response

    return send_file(report_file, as_attachment=True, download_name="psychological_analysis_report.docx")

@app.route('/download', methods=['GET'])
def download():
    output_file = session.get('output_file')
    if not output_file or not os.path.exists(output_file):
        flash("No file available to download.", "error")
        return redirect(url_for('index'))

    if session.get('file_downloaded'):
        flash("The file has already been downloaded.", "error")
        return redirect(url_for('index'))

    session['file_downloaded'] = True

    @after_this_request
    def remove_file(response):
        try:
            os.remove(output_file)
            session.pop('output_file', None)  # Clear session after file is removed
            session.pop('file_downloaded', None)  # Reset download state
            flash("Download complete. The session has been reset.", "success")
        except Exception as e:
            print(f"Error deleting file: {str(e).encode('utf-8', errors='ignore')}")
        return response

    return send_file(output_file, as_attachment=True, download_name="parsed_chat_with_reply_times.xlsx")

@app.route('/report_status/<job_id>', methods=['GET'])
def report_status(job_id):
    if session.get('report_job_id') != job_id:
        return {"error": "Invalid job ID."}, 403

    job = get_report_job(job_id)
    if not job:
        return {"error": "Job not found."}, 404

    if job.get('status') == 'done' and job.get('report_file'):
        session['report_file'] = job['report_file']

    return {
        "job_id": job_id,
        "status": job.get("status"),
        "message": job.get("message", ""),
        "stage": job.get("stage"),
        "chunk_current": job.get("chunk_current"),
        "chunk_total": job.get("chunk_total"),
        "cached": job.get("cached", False)
    }

@app.route('/report_status_page/<job_id>', methods=['GET'])
def report_status_page(job_id):
    if session.get('report_job_id') != job_id:
        flash("Invalid report job.", "error")
        return redirect(url_for('index'))
    return render_template('report_status.html', job_id=job_id)

@app.route('/report_regenerate/<job_id>', methods=['GET'])
def report_regenerate(job_id):
    if session.get('report_job_id') != job_id and job_id not in session.get('report_job_ids', []):
        flash("Invalid report job.", "error")
        return redirect(url_for('index'))
    job = get_report_job(job_id)
    if not job or not job.get('input_file'):
        flash("No source file available for regeneration.", "error")
        return redirect(url_for('index'))
    input_path = job['input_file']
    if not os.path.exists(input_path):
        flash("Source file no longer exists.", "error")
        return redirect(url_for('index'))
    suffix = os.path.splitext(input_path)[1].lower()
    new_job_id = start_report_job(input_path, suffix, use_cache=False)
    return redirect(url_for('report_status_page', job_id=new_job_id))

# Route for processing CSV/Excel files and generating the psychological report
@app.route('/upload_csv', methods=['POST'])
def upload_csv():
    file = request.files.get('file_csv')
    if file and (file.filename.endswith('.csv') or file.filename.endswith('.xlsx')):
        suffix = '.csv' if file.filename.endswith('.csv') else '.xlsx'
        job_id = uuid.uuid4().hex
        input_path = os.path.join(cache_dir, f"report_input_{job_id}{suffix}")
        file.save(input_path)
        start_report_job(input_path, suffix, use_cache=True, job_id=job_id)
        return redirect(url_for('report_status_page', job_id=job_id))

    flash("Invalid file format. Please upload a CSV or Excel file.", 'error')
    return redirect(url_for('index'))

# Function to generate a psychological report
def call_openai_with_backoff(prompt, model_name, min_request_interval, rpm_limit):
    global openai_last_request_ts
    retry_attempts = 5
    max_delay = 60.0

    for attempt in range(retry_attempts):
        now = time.time()
        elapsed = now - openai_last_request_ts
        if elapsed < min_request_interval:
            wait_for = min_request_interval - elapsed
            print(f"Waiting {wait_for:.2f}s to respect RPM limit.")
            time.sleep(wait_for)
        try:
            response = openai.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0
            )
            openai_last_request_ts = time.time()
            return response.choices[0].message.content
        except RateLimitError:
            if rpm_limit < 10:
                delay = 60.0
            else:
                delay = min(
                    max_delay,
                    max(1.0, min_request_interval * (2 ** attempt))
                )
            print(
                "Rate limit exceeded. Retrying in "
                f"{delay:.2f} seconds... (Attempt {attempt + 1}/{retry_attempts})"
            )
            time.sleep(delay)
        except APIError as e:
            return None

    return None


def build_report_metadata(chat_data):
    if chat_data is None or chat_data.empty:
        return ""

    df = chat_data.copy()
    if 'Datetime' in df.columns:
        dt_series = pd.to_datetime(df['Datetime'], errors='coerce')
    elif 'Date' in df.columns and 'Time' in df.columns:
        dt_series = pd.to_datetime(
            df['Date'].astype(str) + ' ' + df['Time'].astype(str),
            errors='coerce'
        )
    elif 'Date' in df.columns:
        dt_series = pd.to_datetime(df['Date'], errors='coerce')
    else:
        dt_series = pd.Series([pd.NaT] * len(df))

    valid_dates = dt_series.dropna().sort_values()
    first_dt = valid_dates.min() if not valid_dates.empty else None
    last_dt = valid_dates.max() if not valid_dates.empty else None

    total_messages = len(df)
    participants = df['Sender'].dropna().astype(str) if 'Sender' in df.columns else pd.Series([], dtype=str)
    participant_count = participants.nunique() if not participants.empty else 0
    top_participants = participants.value_counts().head(3)

    span_days = None
    if first_dt is not None and last_dt is not None:
        span_days = (last_dt - first_dt).days + 1

    longest_gap_hours = None
    if len(valid_dates) > 1:
        gaps = valid_dates.diff().dropna()
        longest_gap = gaps.max()
        longest_gap_hours = longest_gap.total_seconds() / 3600.0

    busiest_day = None
    if 'Date' in df.columns and not df['Date'].isna().all():
        day_counts = df['Date'].astype(str).value_counts()
        if not day_counts.empty:
            busiest_day = (day_counts.index[0], int(day_counts.iloc[0]))

    avg_messages_per_participant = None
    top_participant_share = None
    if participant_count > 0:
        avg_messages_per_participant = total_messages / participant_count
        if not top_participants.empty:
            top_participant_share = (top_participants.iloc[0] / total_messages) * 100

    peak_hour = None
    if 'Time' in df.columns and not df['Time'].isna().all():
        time_values = df['Time'].astype(str).str.strip()
        time_series = pd.to_datetime(time_values, format='%H:%M:%S', errors='coerce')
        if time_series.isna().all():
            time_series = pd.to_datetime(time_values, format='%H:%M', errors='coerce')
        hour_counts = time_series.dt.hour.value_counts().sort_index()
        if not hour_counts.empty:
            peak_hour = (int(hour_counts.idxmax()), int(hour_counts.max()))

    rows = []
    if first_dt is not None:
        rows.append(("First message", first_dt.strftime('%Y-%m-%d %H:%M')))
    if last_dt is not None:
        rows.append(("Last message", last_dt.strftime('%Y-%m-%d %H:%M')))
    rows.append(("Total messages", f"{total_messages}"))
    rows.append(("Participants", f"{participant_count}"))
    if span_days is not None and span_days > 0:
        rows.append(("Time span (days)", f"{span_days}"))
        rows.append(("Messages per day", f"{total_messages / span_days:.1f}"))
    if longest_gap_hours is not None:
        rows.append(("Longest gap (hours)", f"{longest_gap_hours:.1f}"))
    if busiest_day is not None:
        rows.append(("Busiest day", f"{busiest_day[0]} ({busiest_day[1]} messages)"))
    if avg_messages_per_participant is not None:
        rows.append(("Avg messages per participant", f"{avg_messages_per_participant:.1f}"))
    if top_participant_share is not None and not top_participants.empty:
        top_name = str(top_participants.index[0])
        rows.append(("Most active participant", f"{top_name} ({top_participant_share:.1f}%)"))
    if peak_hour is not None:
        rows.append(("Peak hour of day", f"{peak_hour[0]:02d}:00 ({peak_hour[1]} messages)"))
    if not top_participants.empty:
        top_items = ", ".join(f"{name} ({count})" for name, count in top_participants.items())
        rows.append(("Top participants", top_items))

    lines = ["## Report Metadata", "| Metric | Value |", "| --- | --- |"]
    for metric, value in rows:
        lines.append(f"| {metric} | {value} |")

    return "\n".join(lines)


def generate_psychological_report(chat_data, progress_cb=None, use_cache=True):
    # Convert the chat data to a readable format
    chat_text = ""
    for index, row in chat_data.iterrows():
        chat_text += f"[{row['Date']} {row['Time']}] {row['Sender']}: {row['Message']}\n"

    model_name = "gpt-4o-mini"
    rpm_limit = float(os.getenv("OPENAI_RPM_LIMIT", "3"))
    min_request_interval = 60.0 / max(rpm_limit, 1.0)
    chunk_max_chars = int(os.getenv("CHAT_CHUNK_MAX_CHARS", "8000"))
    max_chunks = int(os.getenv("CHAT_MAX_CHUNKS", "50"))

    cache_key_input = f"{model_name}:{chunk_max_chars}:{chat_text}"
    cache_key = hashlib.sha256(cache_key_input.encode('utf-8')).hexdigest()
    cache_path = os.path.join(cache_dir, f"{cache_key}.txt")
    if use_cache and os.path.exists(cache_path):
        with open(cache_path, "r", encoding="utf-8") as cached:
            cached_text = cached.read()
        if progress_cb:
            progress_cb(stage='cache', cached=True, message='Loaded from cache.')
        return cached_text, True

    lines = chat_text.splitlines()
    if max_chunks > 0:
        target_chunk_size = math.ceil(len(chat_text) / max_chunks)
        chunk_max_chars = max(chunk_max_chars, target_chunk_size)
    chunks = []
    current_lines = []
    current_len = 0
    for line in lines:
        line_len = len(line) + 1
        if current_lines and current_len + line_len > chunk_max_chars:
            chunks.append("\n".join(current_lines))
            current_lines = [line]
            current_len = line_len
        else:
            current_lines.append(line)
            current_len += line_len
    if current_lines:
        chunks.append("\n".join(current_lines))

    total_parts = len(chunks)
    if progress_cb:
        progress_cb(stage='chunking', chunk_total=total_parts, chunk_current=0)
    print(f"Prepared {total_parts} chat chunk(s) for report generation.")
    if total_parts > 1 and rpm_limit < 10:
        min_request_interval = 60.0
        print("Chunked mode enabled; waiting 60s before first OpenAI request.")
        time.sleep(min_request_interval)

    metadata_section = build_report_metadata(chat_data)

    if total_parts == 1:
        prompt = PROMPT_TEMPLATE.format(
            part=1,
            total_parts=1,
            chat_text=chunks[0]
        )
        if progress_cb:
            progress_cb(stage='chunks', chunk_total=1, chunk_current=1)
        content = call_openai_with_backoff(
            prompt,
            model_name,
            min_request_interval,
            rpm_limit
        )
        if content is None:
            return "Error: Rate limit exceeded after multiple attempts. Please try again later.", False
        if metadata_section:
            content = f"{metadata_section}\n\n{content}"
        with open(cache_path, "w", encoding="utf-8") as cached:
            cached.write(content)
        return content, False

    partial_summaries = []
    for idx, chunk in enumerate(chunks, start=1):
        if progress_cb:
            progress_cb(stage='chunks', chunk_total=total_parts, chunk_current=idx)
        print(f"Submitting chunk {idx}/{total_parts} to OpenAI.")
        prompt = (
            "You will receive a WhatsApp chat snippet (part {part}/{total_parts}). "
            "Act as an Industrial/Organizational Psychology expert with a specialization in Human "
            "Resource Management and Management. Summarize key psychological signals, emotions, "
            "relationships, and communication patterns found in this snippet. Return a concise "
            "bullet list of findings and notable quotes with speaker names.\n\n"
            "WhatsApp Chat Data:\n"
            "{chat_text}"
        ).format(part=idx, total_parts=total_parts, chat_text=chunk)
        summary = call_openai_with_backoff(
            prompt,
            model_name,
            min_request_interval,
            rpm_limit
        )
        if summary is None:
            return "Error: Rate limit exceeded after multiple attempts. Please try again later.", False
        partial_summaries.append(summary)

    if progress_cb:
        progress_cb(stage='final', chunk_total=total_parts, chunk_current=total_parts)
    print("Submitting final combined summary request to OpenAI.")
    combined_prompt = f"""
    You received {total_parts} partial summaries from a WhatsApp chat. Combine them into a
    single comprehensive psychological analysis report. Focus on emotions, relationships,
    psychological conditions, and communication patterns across the full conversation.

    Partial Summaries:
    {"\n\n".join(partial_summaries)}
    """

    content = call_openai_with_backoff(
        combined_prompt,
        model_name,
        min_request_interval,
        rpm_limit
    )
    if content is None:
        return "Error: Rate limit exceeded after multiple attempts. Please try again later.", False
    if metadata_section:
        content = f"{metadata_section}\n\n{content}"
    with open(cache_path, "w", encoding="utf-8") as cached:
        cached.write(content)
    return content, False


def generate_word_report(analysis_text, file_name):
    # Function to create a Word document with the analysis
    from docx import Document
    doc = Document()
    doc.add_heading('Psychological Analysis Report', 0)

    def is_table_separator(line):
        stripped = line.strip()
        if not stripped.startswith('|') or not stripped.endswith('|'):
            return False
        cells = [cell.strip() for cell in stripped.strip('|').split('|')]
        if not cells:
            return False
        return all(cell and set(cell) <= {'-'} for cell in cells)

    def parse_table_row(line):
        return [cell.strip() for cell in line.strip().strip('|').split('|')]

    def add_markdown_runs(paragraph, text):
        token_pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(token_pattern, text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(part)

    def set_cell_text(cell, text, bold=False):
        cell.text = ""
        para = cell.paragraphs[0]
        if bold:
            run = para.add_run(text)
            run.bold = True
        else:
            add_markdown_runs(para, text)

    def render_markdown_to_docx(text):
        lines = text.splitlines()
        i = 0
        while i < len(lines):
            raw_line = lines[i]
            line = raw_line.rstrip()
            if not line:
                doc.add_paragraph("")
                i += 1
                continue
            if line.startswith('|') and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                header = parse_table_row(line)
                i += 2
                body_rows = []
                while i < len(lines):
                    body_line = lines[i].rstrip()
                    if not body_line or not body_line.strip().startswith('|'):
                        break
                    body_rows.append(parse_table_row(body_line))
                    i += 1
                table = doc.add_table(rows=1, cols=len(header))
                table.style = 'Table Grid'
                for col_idx, value in enumerate(header):
                    set_cell_text(table.cell(0, col_idx), value, bold=True)
                for row in body_rows:
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        if col_idx < len(row_cells):
                            set_cell_text(row_cells[col_idx], value)
                continue
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line[level:].strip()
                doc.add_heading(heading_text, level=min(max(level, 1), 4))
                i += 1
                continue
            if line.startswith(('- ', '* ')):
                para = doc.add_paragraph(style='List Bullet')
                add_markdown_runs(para, line[2:].strip())
                i += 1
                continue
            para = doc.add_paragraph()
            add_markdown_runs(para, line)
            i += 1

    render_markdown_to_docx(analysis_text)

    # Save the document as a temporary file
    word_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(word_file.name)

    return word_file.name


def generate_html_report(analysis_text, file_name):
    def escape_text(text):
        return html.escape(text, quote=True)

    def add_inline_md(text):
        token_pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(token_pattern, text)
        rendered = []
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                rendered.append(f"<strong>{escape_text(part[2:-2])}</strong>")
            elif part.startswith('*') and part.endswith('*'):
                rendered.append(f"<em>{escape_text(part[1:-1])}</em>")
            elif part.startswith('`') and part.endswith('`'):
                rendered.append(f"<code>{escape_text(part[1:-1])}</code>")
            else:
                rendered.append(escape_text(part))
        return "".join(rendered)

    def is_table_separator(line):
        stripped = line.strip()
        if not stripped.startswith('|') or not stripped.endswith('|'):
            return False
        cells = [cell.strip() for cell in stripped.strip('|').split('|')]
        if not cells:
            return False
        return all(cell and set(cell) <= {'-'} for cell in cells)

    def parse_table_row(line):
        return [cell.strip() for cell in line.strip().strip('|').split('|')]

    lines = analysis_text.splitlines()
    body_parts = ["<div class=\"report\">"]
    i = 0
    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.rstrip()
        check_line = line.lstrip()
        if not line:
            body_parts.append("<p></p>")
            i += 1
            continue
        if check_line.startswith('|') and i + 1 < len(lines) and is_table_separator(lines[i + 1].lstrip()):
            header = parse_table_row(check_line)
            i += 2
            body_rows = []
            while i < len(lines):
                body_line = lines[i].rstrip()
                body_check = body_line.lstrip()
                if not body_line or not body_check.startswith('|'):
                    break
                body_rows.append(parse_table_row(body_check))
                i += 1
            body_parts.append("<table class=\"report-table\">")
            body_parts.append("<thead><tr>")
            for cell in header:
                body_parts.append(f"<th>{add_inline_md(cell)}</th>")
            body_parts.append("</tr></thead><tbody>")
            for row in body_rows:
                body_parts.append("<tr>")
                for cell in row:
                    body_parts.append(f"<td>{add_inline_md(cell)}</td>")
                body_parts.append("</tr>")
            body_parts.append("</tbody></table>")
            continue
        if check_line.startswith('#'):
            level = len(check_line) - len(check_line.lstrip('#'))
            heading_text = check_line[level:].strip()
            level = max(1, min(level, 4))
            body_parts.append(f"<h{level}>{add_inline_md(heading_text)}</h{level}>")
            i += 1
            continue
        if check_line.startswith(('- ', '* ')):
            items = []
            while i < len(lines):
                item_line = lines[i].lstrip()
                if not item_line.startswith(('- ', '* ')):
                    break
                items.append(item_line[2:].strip())
                i += 1
            body_parts.append("<ul>")
            for item in items:
                body_parts.append(f"<li>{add_inline_md(item)}</li>")
            body_parts.append("</ul>")
            continue
        body_parts.append(f"<p>{add_inline_md(check_line)}</p>")
        i += 1

    body_parts.append("</div>")
    html_doc = "\n".join([
        "<!DOCTYPE html>",
        "<html lang=\"en\">",
        "<head>",
        "<meta charset=\"UTF-8\">",
        "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">",
        "<title>Psychological Report</title>",
        "<style>",
        "body{font-family:Arial, sans-serif; margin:24px; color:#111;}",
        ".report-table{border-collapse:collapse; width:100%; margin:16px 0;}",
        ".report-table th,.report-table td{border:1px solid #ccc; padding:8px;}",
        ".report-table th{background:#f2f2f2; text-align:left;}",
        "code{background:#f6f6f6; padding:2px 4px; border-radius:4px;}",
        "</style>",
        "</head>",
        "<body>",
        "\n".join(body_parts),
        "</body>",
        "</html>"
    ])

    html_file = tempfile.NamedTemporaryFile(delete=False, suffix=".html")
    with open(html_file.name, "w", encoding="utf-8") as handle:
        handle.write(html_doc)
    return html_file.name


@app.route('/reports', methods=['GET'])
def report_list():
    job_ids = session.get('report_job_ids', [])
    reports = []
    for job_id in job_ids:
        job = get_report_job(job_id)
        if not job or job.get('status') != 'done':
            continue
        report_file = job.get('report_file')
        html_file = job.get('html_file')
        if report_file and os.path.exists(report_file):
            reports.append({
                "job_id": job_id,
                "report_file": report_file,
                "html_file": html_file
            })
    return render_template('report_list.html', reports=reports)


@app.route('/report_html/<job_id>', methods=['GET'])
def report_html(job_id):
    if session.get('report_job_id') != job_id and job_id not in session.get('report_job_ids', []):
        flash("Invalid report job.", "error")
        return redirect(url_for('index'))
    job = get_report_job(job_id)
    if not job or not job.get('html_file') or not os.path.exists(job['html_file']):
        flash("No HTML report available.", "error")
        return redirect(url_for('index'))
    return send_file(job['html_file'])

@app.route('/report_docx/<job_id>', methods=['GET'])
def report_docx(job_id):
    if session.get('report_job_id') != job_id and job_id not in session.get('report_job_ids', []):
        flash("Invalid report job.", "error")
        return redirect(url_for('index'))
    job = get_report_job(job_id)
    if not job or not job.get('report_file') or not os.path.exists(job['report_file']):
        flash("No report available.", "error")
        return redirect(url_for('index'))

    report_file = job['report_file']

    @after_this_request
    def remove_report_file(response):
        try:
            os.remove(report_file)
            job['report_file'] = None
        except Exception as e:
            print(f"Error deleting file: {str(e)}")
        return response

    return send_file(report_file, as_attachment=True, download_name="psychological_analysis_report.docx")

if __name__ == '__main__':
    host = os.getenv('FLASK_RUN_HOST', '127.0.0.1')
    port = int(os.getenv('FLASK_RUN_PORT', '5000'))
    app.run(debug=True, host=host, port=port)
